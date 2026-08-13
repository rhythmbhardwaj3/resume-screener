"""
parser.py

Extracts plain text from resumes / job descriptions in PDF, DOCX, or TXT format.
"""
import os


def extract_text_from_pdf(filepath_or_bytes):
    import pdfplumber
    text_chunks = []
    with pdfplumber.open(filepath_or_bytes) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_chunks.append(page_text)
    return "\n".join(text_chunks)


def extract_text_from_docx(filepath_or_bytes):
    import docx
    document = docx.Document(filepath_or_bytes)
    paragraphs = [p.text for p in document.paragraphs]
    # also grab text inside tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def extract_text_from_txt(filepath_or_bytes):
    if hasattr(filepath_or_bytes, "read"):
        raw = filepath_or_bytes.read()
        if isinstance(raw, bytes):
            return raw.decode("utf-8", errors="ignore")
        return raw
    with open(filepath_or_bytes, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_text(file_path_or_obj, filename=None):
    """
    Dispatches to the correct extractor based on file extension.
    file_path_or_obj: a filesystem path (str) OR a file-like object (e.g. Streamlit UploadedFile)
    filename: required if passing a file-like object, to determine extension
    """
    name = filename if filename else file_path_or_obj
    ext = os.path.splitext(str(name))[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path_or_obj)
    elif ext == ".docx":
        return extract_text_from_docx(file_path_or_obj)
    elif ext in (".txt", ".text"):
        return extract_text_from_txt(file_path_or_obj)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: .pdf, .docx, .txt")
